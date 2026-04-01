from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT_NAME = "Arial"

doc = Document()

# ── Page setup: US Letter, 1-inch margins ──────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Default style ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Pt(11)
style.font.color.rgb = BLACK

# ── Footer ─────────────────────────────────────────────────────────────────
def add_footer(sec):
    footer = sec.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Bryce Mack | Shift4 SkyTab POS    ")
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    # page number field
    fld_xml = (
        '<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>'
    )
    run2 = para.add_run()
    run2.font.name = FONT_NAME
    run2.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run2._r.append(fld_begin)

    run3 = para.add_run()
    run3.font.name = FONT_NAME
    run3.font.size = Pt(9)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run3._r.append(instr)

    run4 = para.add_run()
    run4.font.name = FONT_NAME
    run4.font.size = Pt(9)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run4._r.append(fld_end)

add_footer(section)

# ── Helper: plain paragraph ────────────────────────────────────────────────
def add_plain(text="", bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    # zero space before/after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if text:
        run = p.add_run(text)
        run.font.name  = FONT_NAME
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.color.rgb = color if color else BLACK
    return p

# ── Helper: blank line ─────────────────────────────────────────────────────
def blank():
    return add_plain()

# ── Helper: section header ─────────────────────────────────────────────────
def section_header(text):
    p = add_plain(text, bold=True, size=14, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p

# ── Helper: template label ─────────────────────────────────────────────────
def template_label(text):
    p = add_plain(text, bold=True, size=12, color=NAVY)
    return p

# ── Helper: subject line ───────────────────────────────────────────────────
def subject_line(value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    bold_run = p.add_run("Subject: ")
    bold_run.font.name  = FONT_NAME
    bold_run.font.size  = Pt(11)
    bold_run.font.bold  = True
    bold_run.font.color.rgb = BLACK
    val_run = p.add_run(value)
    val_run.font.name  = FONT_NAME
    val_run.font.size  = Pt(11)
    val_run.font.bold  = False
    val_run.font.color.rgb = BLACK
    return p

# ── Helper: horizontal rule (paragraph bottom border) ─────────────────────
def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B3A6B")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Helper: body paragraph ─────────────────────────────────────────────────
def body(text):
    return add_plain(text)

# ── Helper: double-space gap between templates ─────────────────────────────
def gap():
    blank()
    blank()

# ══════════════════════════════════════════════════════════════════════════
# CONTENT
# ══════════════════════════════════════════════════════════════════════════

# ── VOICEMAIL & CALL SCRIPTS ──────────────────────────────────────────────
section_header("VOICEMAIL & CALL SCRIPTS")
blank()

template_label("VOICEMAIL SCRIPT")
blank()
body("Good morning/afternoon. This message is for [Name]. This is [Your Name] with SkyTab point of sales system. You had requested information. You can reach me at [Phone Number] extension [Extension]. I'm also going to send you an email and a text message, so feel free to reply and let me know a good time for you to connect so we can go over the program. Thank you and have a great day.")
blank()
add_divider()
gap()

# ── EMAIL TEMPLATES ───────────────────────────────────────────────────────
section_header("EMAIL TEMPLATES")
blank()

# E-1
template_label("E-1: INITIAL OUTREACH — DAY 1")
subject_line("Your SkyTab POS Info — Let's Connect, [First Name]")
blank()
body("Hi [First Name],")
blank()
body("Thanks for reaching out about SkyTab POS! I'm [Your Name] with Shift4 and I'd love to help get your questions answered.")
blank()
body("SkyTab is one of the most affordable, feature-rich POS systems built specifically for restaurants — with $0 upfront hardware, free professional installation, and built-in tools for online ordering, loyalty programs, and more.")
blank()
body("Do you have 10–15 minutes this week for a quick call or Google Meet? I can walk you through everything and put together a custom quote for your restaurant.")
blank()
body("What's your availability?")
blank()
body("Best,")
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-2
template_label("E-2: VALUE PROPOSITION — DAY 2")
subject_line("What Makes SkyTab Different, [First Name]")
blank()
body("Hi [First Name],")
blank()
body("Tried to reach you yesterday — just wanted to make sure this didn't fall through the cracks.")
blank()
body("Here's why hundreds of restaurant owners choose SkyTab every month:")
blank()
body("$0 upfront — all hardware on monthly subscription with lifetime warranty")
body("Free professional installation + menu programming (we build it for you)")
body("Commission-free online ordering")
body("Built-in loyalty program, reservations, inventory, employee scheduling")
body("24/7 U.S.-based support")
body("Advantage Program — offset or eliminate credit card processing fees entirely")
blank()
body("Most restaurants we work with save thousands per year on processing costs alone.")
blank()
body("Worth a quick 15-minute conversation? Just reply or call me at [Phone Number].")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-3
template_label("E-3: SOCIAL PROOF — DAY 3")
subject_line("What Restaurant Owners Are Saying About SkyTab")
blank()
body("Hi [First Name],")
blank()
body("Quick follow-up — I know running a restaurant keeps you slammed, so I'll keep this short.")
blank()
body("SkyTab was built from the ground up for food and beverage operations. The results we see consistently:")
blank()
body("Faster table turns with tableside ordering and payment on handhelds")
body("Higher check averages from built-in upselling prompts")
body("Better labor control with integrated scheduling and time clock")
body("More repeat customers through the loyalty and email marketing tools")
blank()
body("And with our Advantage Program, your credit card processing fees go from thousands per year to virtually zero.")
blank()
body("I'd love to show you the system live — takes about 15 minutes on Google Meet. Are you free [Day] or [Day] this week?")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-4
template_label("E-4: PAIN POINT — DAY 5")
subject_line("Are Processing Fees Eating Into Your Margins, [First Name]?")
blank()
body("Hi [First Name],")
blank()
body("Restaurant margins are tight — that's no secret. Most operators are losing 2.5–3.5% of every credit card transaction to processing fees, which adds up fast.")
blank()
body("With SkyTab's Advantage Program, you can bring that cost to virtually zero through dual pricing — cash customers pay your standard price, credit card customers see a slightly adjusted price that covers the processing cost. Your business keeps 100% of every sale.")
blank()
body("On top of that, you'd be getting a full enterprise-level POS system with:")
blank()
body("Seamless 3rd-party delivery integration (DoorDash, Uber Eats, Grubhub — no extra tablets)")
body("Cloud reporting you can check from your phone")
body("Free menu build, free installation, lifetime warranty on hardware")
blank()
body("Would love to put together a custom quote for [Restaurant Name]. Can we find 15 minutes this week?")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-5
template_label("E-5: COMPETITION ANGLE — DAY 7")
subject_line("SkyTab vs. Toast/Square/Clover — Quick Comparison")
blank()
body("Hi [First Name],")
blank()
body("I've reached out a few times and haven't been able to connect — I'll keep this brief.")
blank()
body("If you're comparing options, here's how SkyTab stacks up:")
blank()
body("Upfront hardware cost: SkyTab $0 | Toast $500–$3,000+ | Square $49–$799 | Clover $500–$1,500+")
body("Monthly software: SkyTab $20 | Toast $69–$165+ | Square $60+ | Clover $39–$290+")
body("Processing fee offset: SkyTab Yes (Advantage Program) | Toast Limited | Square No | Clover No")
body("Free installation: SkyTab Yes | Toast Paid | Square Self-install | Clover Varies")
body("24/7 U.S.-based support: SkyTab Yes | Toast No | Square No | Clover No")
body("Lifetime hardware warranty: SkyTab Yes | Toast No | Square No | Clover No")
blank()
body("SkyTab is purpose-built for restaurants, backed by Shift4 — one of the largest payment processors in the U.S.")
blank()
body("Still worth a conversation? Reply here or call me at [Phone Number]. Takes 15 minutes.")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-6
template_label("E-6: BREAK-UP EMAIL — DAY 10")
subject_line("Closing the Loop, [First Name]")
blank()
body("Hi [First Name],")
blank()
body("I've reached out several times and haven't been able to connect, so I wanted to send one final note.")
blank()
body("If the timing isn't right or you've gone a different direction — no worries at all. I'll get out of your inbox.")
blank()
body("But if you're still exploring options and just haven't had a chance to respond, I'm happy to work around your schedule. Even a 10-minute call could save you thousands this year.")
blank()
body("Either way, feel free to reach out anytime — my info is below and I'm always happy to help.")
blank()
body("Wishing you success with [Restaurant Name],")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-7
template_label("E-7: DEMO CONFIRMATION — AFTER BOOKING")
subject_line("We're Confirmed — [Day/Date] at [Time]")
blank()
body("Hi [First Name],")
blank()
body("Looking forward to our Google Meet on [Day], [Date] at [Time]!")
blank()
body("Here's the link: [Google Meet Link]")
blank()
body("In our 15–20 minutes together I'll cover:")
blank()
body("A live walkthrough of the SkyTab POS interface")
body("The Advantage Program and how it can eliminate your processing fees")
body("Pricing and what's included (hardware, installation, support)")
body("A custom quote for your setup")
blank()
body("If anything comes up and you need to reschedule, just let me know — totally flexible.")
blank()
body("See you then!")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-8
template_label("E-8: POST-DEMO RECAP — WITHIN 1 HOUR OF DEMO")
subject_line("Your SkyTab Proposal — [Restaurant Name]")
blank()
body("Hi [First Name],")
blank()
body("Great connecting with you today! As promised, here's a summary of what we discussed along with your custom quote.")
blank()
body("[ATTACH PROPOSAL HERE]")
blank()
body("Key highlights for your setup:")
body("[List 2–3 specifics from your conversation]")
body("Hardware at $0 upfront with lifetime warranty")
body("Lighthouse software at $20/month (first 2 months free)")
body("Advantage Program to bring your processing costs to near zero")
body("First year main system fee waived with Advantage enrollment — $49.99/month year one")
blank()
body("Next steps are simple — the application takes about 10 minutes over the phone, approval usually comes within 24 hours, and once approved your launch control rep will build and program the entire system before it ships.")
blank()
body("Any questions? Happy to jump back on a call whenever works for you.")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# E-9
template_label("E-9: POST-CLOSE WELCOME — AFTER APPLICATION SUBMITTED")
subject_line("Welcome to SkyTab, [First Name]! Here's What's Next")
blank()
body("Hi [First Name],")
blank()
body("Congratulations and welcome to the SkyTab family! We're genuinely excited to be part of [Restaurant Name]'s journey.")
blank()
body("Here's what to expect next:")
blank()
body('1. Watch for your "Welcome to Shift4" email — click the blue Register button, create your login, and answer the 4–5 business questions. If your menu isn\'t ready yet, just attach any photo as a placeholder.')
blank()
body("2. Send me a voided check (write VOID on a check and email me a photo). If you don't have checks yet, your bank can provide a letter on letterhead with your name, business name, account number, routing number, and a bank rep signature.")
blank()
body("3. Once your account is approved and documentation is received, you'll be assigned a dedicated launch control rep who will build and program your entire system.")
blank()
body("4. A professional installer will coordinate with you to schedule setup and train your staff.")
blank()
body("In the meantime, I'm your main point of contact for anything you need. Treat me as home base — call or text me anytime.")
blank()
body("Can't wait to see [Restaurant Name] thrive!")
blank()
body("[Your Name]")
body("Shift4 | SkyTab POS")
body("[Phone Number]")
blank()
add_divider()
gap()

# ── SMS TEMPLATES ─────────────────────────────────────────────────────────
section_header("SMS TEMPLATES")
blank()

# SMS-1
template_label("SMS-1: INITIAL TEXT — DAY 1 (immediately after first call attempt)")
blank()
body("Hi [First Name], this is [Your Name] with SkyTab POS — I just tried calling. You reached out about our POS system and I'd love to help. Do you have a few minutes today or tomorrow?")
blank()
add_divider()
gap()

# SMS-2
template_label("SMS-2: DAY 2 FOLLOW-UP")
blank()
body("Hey [First Name], [Your Name] from SkyTab again. Wanted to make sure you got my info. We've got a promo running — main system free for 12 months with our Advantage Program. Worth a quick 15-min call. When's a good time?")
blank()
add_divider()
gap()

# SMS-3
template_label("SMS-3: DAY 5 CHECK-IN")
blank()
body("Hi [First Name] — [Your Name] with SkyTab. Still happy to help if the timing's right. Even a 10-min call could save you thousands in processing fees this year. Totally low pressure.")
blank()
add_divider()
gap()

# SMS-4
template_label("SMS-4: FINAL CHECK-IN — DAY 10")
blank()
body("Hey [First Name], last reach out from [Your Name] at SkyTab. If you're still looking for a POS system or want to revisit down the road, I'm here. Wishing you the best either way!")
blank()
add_divider()
gap()

# SMS-5
template_label("SMS-5: DEMO REMINDER — NIGHT BEFORE")
blank()
body("Hi [First Name]! Just a reminder about our Google Meet tomorrow at [Time]. Here's the link: [Link]. Looking forward to it! — [Your Name], SkyTab")
blank()
add_divider()
gap()

# SMS-6
template_label("SMS-6: POST-DEMO FOLLOW-UP — SAME DAY AS DEMO")
blank()
body("Hey [First Name], great chatting today! Just sent your custom quote to [Email]. Let me know if you have any questions or want to move forward — process takes about 10 min! — [Your Name]")
blank()
add_divider()

# ── Save ──────────────────────────────────────────────────────────────────
out_path = r"C:\Users\bmack\Desktop\Inside Sales\Templates - Copy & Paste (Claude).docx"
doc.save(out_path)
print(f"Saved: {out_path}")
