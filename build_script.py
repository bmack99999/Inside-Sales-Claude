# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:/Users/bmack/Desktop/Inside Sales/Farzad Method - Sales Script (Claude).docx"

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_para_bg(para, color):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(color))
    pPr.append(shd)


def add_section_header(doc, text):
    para = doc.add_paragraph()
    set_para_bg(para, NAVY)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    run.font.name = "Arial"
    return para


def add_subheader(doc, text):
    para = doc.add_paragraph()
    set_para_bg(para, LIGHT_BLUE)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.font.name = "Arial"
    return para


def bullet(doc, text, indent=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Inches(0.5 * (indent + 1))
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = BLACK
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = BLACK
    return para


def normal(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = BLACK
    return para


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    r = para.add_run("Bryce Mack | Shift4 SkyTab POS     Page ")
    r.font.size = Pt(9)
    r.font.name = "Arial"
    # PAGE field
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.text = "PAGE"
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    r2 = para.add_run(); r2.font.size = Pt(9); r2.font.name = "Arial"
    r2._r.append(fc1); r2._r.append(ins); r2._r.append(fc2)
    r3 = para.add_run(" of "); r3.font.size = Pt(9); r3.font.name = "Arial"
    # NUMPAGES field
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "begin")
    ins2 = OxmlElement("w:instrText"); ins2.text = "NUMPAGES"
    fc4 = OxmlElement("w:fldChar"); fc4.set(qn("w:fldCharType"), "end")
    r4 = para.add_run(); r4.font.size = Pt(9); r4.font.name = "Arial"
    r4._r.append(fc3); r4._r.append(ins2); r4._r.append(fc4)


# ── BUILD ──────────────────────────────────────────────────────────────────────
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)

add_footer(doc)

# ── TITLE BLOCK ────────────────────────────────────────────────────────────────
t1 = doc.add_paragraph()
set_para_bg(t1, NAVY)
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1.paragraph_format.space_before = Pt(6)
t1.paragraph_format.space_after = Pt(0)
r = t1.add_run("Farzad Method \u2014 SkyTab POS Sales Script")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = WHITE; r.font.name = "Arial"

t2 = doc.add_paragraph()
set_para_bg(t2, NAVY)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.paragraph_format.space_before = Pt(0)
t2.paragraph_format.space_after = Pt(10)
r = t2.add_run("Bryce Mack | Shift4 Payments")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE; r.font.name = "Arial"

# ── SECTION 1 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 1: VOICEMAIL SCRIPT")
bullet(doc, "\u201cGood morning/afternoon. This message is for [Name]. This is [Your Name] with SkyTab point of sales system. You had requested information. You can reach me at [number] extension [ext]. I\u2019m also going to send you an email and a text message. So feel free to reply to me and let me know a good time for you to connect so we can go over the program. Thank you and have a great day.\u201d")

# ── SECTION 2 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 2: OPENING")
bullet(doc, "\u201cHi [Name], this is [Your Name] with SkyTab point of sales system. How are you?\u201d")
bullet(doc, "\u201cI understand you were looking for some information about SkyTab \u2014 I was just reaching back out.\u201d")
bullet(doc, "\u201cIs it a good time or did I catch you at a bad time?\u201d")
bullet(doc, "[If bad time]: \u201cNo problem at all \u2014 when would be a better time for me to reach you?\u201d")
bullet(doc, "[If good time]: \u201cPerfect. Let\u2019s do this.\u201d")

# ── SECTION 3 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 3: DISCOVERY")
bullet(doc, "\u201cBefore I go into SkyTab and the program and all the good stuff that comes with it, I would love to learn a little bit about what you do \u2014 just a little background on your business so we can figure out what\u2019s going to work best for you.\u201d")
bullet(doc, "\u201cWhat kind of business is it? Restaurant, bar, food truck?\u201d")
bullet(doc, "\u201cAre you existing or is it a brand new business?\u201d")
bullet(doc, "[If new]: \u201cAwesome \u2014 exciting! When are you guys opening up?\u201d")
bullet(doc, "[If existing]: \u201cGreat \u2014 how long have you been open? What system are you currently using?\u201d")
bullet(doc, "\u201cHow many devices or stations do you think you\u2019ll need?\u201d")
bullet(doc, "\u201cWill you be doing any tableside ordering \u2014 servers taking orders at the table?\u201d")
bullet(doc, "\u201cAny interest in online ordering, delivery \u2014 DoorDash, Uber Eats, that kind of thing?\u201d")
bullet(doc, "\u201cAnd just based on what you\u2019ve told me\u2026\u201d [note anything specific to recommend \u2014 e.g. customer-facing screen for walk-up windows, handhelds for full-service]")

# ── SECTION 4 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 4: WHO WE ARE \u2014 CREDIBILITY")
bullet(doc, "\u201cJust so you know \u2014 are you familiar with SkyTab at all?\u201d")
bullet(doc, "\u201cSkyTab is a restaurant-built POS platform and it\u2019s powered by Shift4.\u201d")
bullet(doc, "\u201cShift4 is our main company. We\u2019ve been doing POS systems for the last almost 30 years.\u201d")
bullet(doc, "\u201cOur systems have always been restaurant specific \u2014 bar, entertainment, and hospitality specific. Built for a restaurant, unlike Clover which is built for retail.\u201d")
bullet(doc, "\u201cOur technology powers some of the biggest names in food, hospitality, and entertainment:\u201d")
bullet(doc, "Outback Steakhouse", indent=1)
bullet(doc, "Denny\u2019s", indent=1)
bullet(doc, "Dairy Queen", indent=1)
bullet(doc, "Kentucky Fried Chicken", indent=1)
bullet(doc, "\u201cWe also handle transactions for about 75% of all major US stadiums.\u201d")
bullet(doc, "\u201cWe are fully integrated in the San Francisco 49ers stadium \u2014 which just held the Super Bowl. Our SkyTab reporting system told us that they processed $2.3 million in credit cards before halftime through our SkyTab system.\u201d")
bullet(doc, "\u201cWe support hotel groups and have a presence in roughly 50% of the Las Vegas Strip \u2014 restaurants, nightclubs, bars, and hotels.\u201d")
bullet(doc, "\u201cSo whether you\u2019re running a food truck, opening up your first business, or managing multiple locations \u2014 you\u2019re using that same enterprise-grade technology trusted by these global brands, stadiums, and hotels.\u201d")

# ── SECTION 5 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 5: PRICING & WHAT\u2019S INCLUDED")
add_subheader(doc, "A. HARDWARE")
bullet(doc, "\u201cNow \u2014 enough about us. The best part of our program is that it\u2019s free upfront. Zero cost upfront. And it\u2019s $29.99 per month for each device, depending on which devices you need.\u201d")
bullet(doc, "\u201cLet\u2019s pause for a minute and take a look at that.\u201d")
bullet(doc, "\u201cOur main system \u2014 the main hub \u2014 is a bundle: printer, cash drawer, and monitor. How many main systems would you need?\u201d")
bullet(doc, "\u201cWe also have:\u201d")
bullet(doc, "Handheld tablets \u2014 $29.99/month (tableside ordering and payment)", indent=1)
bullet(doc, "Customer-facing screen \u2014 $29.99/month (customer inserts their own card)", indent=1)
bullet(doc, "Kitchen printer \u2014 $10/month additional (can always add later)", indent=1)
bullet(doc, "Kitchen display screen \u2014 available if needed", indent=1)
bullet(doc, "\u201cOur monthly fee includes a full lifetime warranty.\u201d")
bullet(doc, "\u201cIn a restaurant \u2014 fast-paced, heavy traffic, screens getting hit, greasy \u2014 systems can break down. If anything breaks or stops working, we send you new equipment. We can even overnight it. That\u2019s included in the monthly fee.\u201d")

add_subheader(doc, "B. SOFTWARE \u2014 $20/MONTH")
bullet(doc, "\u201cThe next part of our program is our software fee \u2014 $20 a month. It comes with a lot of additional services. Let me list those for you right now so you know exactly what you\u2019re getting.\u201d")
bullet(doc, "We build and configure the system before we ship it \u2014 menu, modifiers, drinks, everything")
bullet(doc, "Free professional SkyTab installer comes to your location \u2014 no charge for installation")
bullet(doc, "Installer provides on-site training \u2014 invite your managers and staff")
bullet(doc, "DoorDash, GrubHub, Uber Eats, Postmates \u2014 already integrated, no third-party needed")
bullet(doc, "Loyalty program \u2014 point system to reward regulars")
bullet(doc, "Free professional restaurant website with online ordering built in")
bullet(doc, "Gift card program")
bullet(doc, "Inventory management \u2014 low stock alerts built in")
bullet(doc, "Employee time clock and labor tracking \u2014 print reports, easy payroll")
bullet(doc, "Cloud reporting and real-time analytics \u2014 access from smartphone app or website from anywhere")
bullet(doc, "24/7 U.S.-based customer service \u2014 no language barriers, questions resolved immediately")

# ── SECTION 6 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 6: ADVANTAGE PROGRAM")
bullet(doc, "\u201cAs far as credit card processing fees go \u2014 we also include our Advantage Program.\u201d")
bullet(doc, "\u201cThis program is very popular right now. It can bring your credit card processing costs down to virtually zero.\u201d")
bullet(doc, "\u201cWe use a program called dual pricing \u2014 it\u2019s built directly into SkyTab. You can legally and transparently pass through a small fee to your customer and keep more of your revenue.\u201d")
bullet(doc, "\u201cIt\u2019s fully compliant with Visa, Mastercard, and the federal government.\u201d")
bullet(doc, "\u201c99% of our customers go with this program. You save thousands of dollars a year because you\u2019re not paying Visa, Mastercard, all those credit card fees.\u201d")
bullet(doc, "\u201cThere is a promotion attached to it \u2014 you get your main system free for the first year when you go with the Advantage Program.\u201d")
bullet(doc, "\u201cSo the $29.99 for the main system is waived for the first 12 months.\u201d")
bullet(doc, "\u201cExample: if you need one main system and one customer-facing screen \u2014 in year one you\u2019d pay $20 software + $29.99 for the screen = $49.99/month. No credit card fees on top of that.\u201d")
bullet(doc, "\u201cAnd I\u2019m always very transparent \u2014 in year two, that $29.99 for the main system comes back on. Just want to make sure you know that upfront.\u201d")

# ── SECTION 7 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 7: CLOSE \u2014 APPLICATION")
bullet(doc, "\u201cIt takes about 10 minutes to submit the application \u2014 very simple. We just knock it out over the phone.\u201d")
bullet(doc, "\u201cOnce we submit, it gets approved within about 24 hours or less.\u201d")
bullet(doc, "\u201cOnce approved, one of our launch control representatives reaches out to you \u2014 you\u2019ll have a specific person assigned to work with you one-on-one.\u201d")
bullet(doc, "\u201cThey\u2019ll build the full system out, present it to you, and once everything looks good they ship it.\u201d")
bullet(doc, "\u201cThen they\u2019ll work with you to schedule a day for our professional installer to come to your location.\u201d")
bullet(doc, "\u201cDo you have about 10 minutes? We can get started right now.\u201d")

# ── SECTION 8 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 8: APPLICATION \u2014 INFO TO COLLECT")
numbered(doc, "Business legal name (and LLC/Inc if applicable)")
numbered(doc, "EIN / Tax ID number")
numbered(doc, "Best email address")
numbered(doc, "Phone number (business or cell)")
numbered(doc, "Business address")
numbered(doc, "Home address (if different \u2014 used for LLC paperwork / device shipping)")
numbered(doc, "Online ordering plans (DoorDash, Uber Eats, website?)")
numbered(doc, "Estimated monthly credit card volume (big month)")
numbered(doc, "Average ticket sale per customer")
numbered(doc, "Owner first name, last name, date of birth")
numbered(doc, "Social Security number (\u201cWe run what\u2019s called an OFAC report \u2014 it\u2019s not a hard pull on your credit. It just checks if you\u2019ve had an account like this closed due to fraud in the past.\u201d)")
numbered(doc, "Are you 100% owner? Any other owners?")
numbered(doc, "Authorized contact person \u2014 name and phone number")
numbered(doc, "Bank name, routing number, account number")
numbered(doc, "Device shipping address (home vs. business \u2014 ask about mailbox security)")
numbered(doc, "Business hours (latest close time)")

normal(doc, "Then: Send the Shift4 application link via email. Walk them through:")
bullet(doc, "Click the blue \"Get Started\" button")
bullet(doc, "Review info on next page \u2014 scroll to bottom")
bullet(doc, "Click blue \"Proceed to Signing\"")
bullet(doc, "Scroll down to find the blue \"I Agree\" button (don\u2019t do physical signature \u2014 it\u2019ll look like a PDF, scroll past it)")

# ── SECTION 9 ─────────────────────────────────────────────────────────────────
add_section_header(doc, "SECTION 9: POST-APPLICATION \u2014 NEXT STEPS TO TELL THE CUSTOMER")
numbered(doc, "\u201cWe\u2019ll get approved within the next 24 hours.\u201d")
numbered(doc, "\u201cYou\u2019ll get a \u2018Welcome to Shift4\u2019 email today \u2014 open it and click the blue Register button.\u201d")
numbered(doc, "\u201cCreate your login with your email and a password.\u201d")
numbered(doc, "\u201cYou\u2019ll be asked 4\u20135 questions about the business \u2014 answer those and attach your menu. If you don\u2019t have the menu yet, just attach any photo \u2014 a picture of anything. We just need those questions completed to get your launch control rep assigned.\u201d")
numbered(doc, "\u201cYour launch control rep will reach out and you\u2019ll work with them one-on-one from there.\u201d")
numbered(doc, "\u201cWe\u2019ll also need a voided check \u2014 just write VOID on a check and email me a photo.\u201d")
numbered(doc, "\u201cIf you don\u2019t have checks yet, you can get a bank letter from your bank on their letterhead. It needs to include: your name, business name, account number, routing number, and a signature from a bank representative.\u201d")
numbered(doc, "\u201cJust treat me as home base \u2014 call or email me anytime if you need anything.\u201d")

doc.save(OUTPUT_PATH)
print("SAVED OK")
